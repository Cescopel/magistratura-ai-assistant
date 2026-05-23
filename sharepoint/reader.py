import os
from config.settings import DOCUMENTS_PATH
from docx import Document

# ============================
# CONFIGURAZIONE CHUNKING
# ============================

CHUNK_SIZE = 30  # Paragrafi per chunk (~1-2 pagine)
CHUNK_OVERLAP = 5  # Sovrapposizione tra chunk

# ============================
# FUNZIONI BASE (come prima)
# ============================

def list_all_documents():
    """
    Elenca tutti i documenti disponibili.
    """
    try:
        documents = []
        
        folders_to_check = [
            DOCUMENTS_PATH,
            "documents_uploaded"
        ]
        
        for folder in folders_to_check:
            if not os.path.exists(folder):
                continue
            
            for filename in os.listdir(folder):
                filepath = os.path.join(folder, filename)
                
                if os.path.isfile(filepath):
                    file_size = os.path.getsize(filepath)
                    documents.append({
                        "name": filename,
                        "path": filepath,
                        "size": file_size
                    })
        
        print(f"📚 Trovati {len(documents)} documenti")
        return documents
    
    except Exception as e:
        print(f"❌ Errore lettura documenti: {e}")
        return []


def search_documents(query):
    """
    Cerca documenti per nome file con fuzzy matching.
    """
    try:
        from fuzzywuzzy import fuzz
        
        all_docs = list_all_documents()
        
        # Prova prima match esatto (veloce)
        exact_matches = [
            doc for doc in all_docs 
            if query.lower() in doc["name"].lower()
        ]
        
        if exact_matches:
            print(f"🔍 Trovati {len(exact_matches)} match esatti per '{query}'")
            return exact_matches
        
        # Se nessun match esatto, prova fuzzy
        print(f"⚠️ Nessun match esatto per '{query}', provo fuzzy matching...")
        
        fuzzy_results = []
        for doc in all_docs:
            # Calcola similarità (0-100)
            score = fuzz.partial_ratio(query.lower(), doc["name"].lower())
            
            if score > 60:  # Soglia 60%
                doc_copy = doc.copy()
                doc_copy["fuzzy_score"] = score
                fuzzy_results.append(doc_copy)
        
        # Ordina per score (migliori prima)
        fuzzy_results.sort(key=lambda x: x["fuzzy_score"], reverse=True)
        
        print(f"🔍 Trovati {len(fuzzy_results)} match fuzzy per '{query}'")
        if fuzzy_results:
            print(f"   Top match: '{fuzzy_results[0]['name']}' (score: {fuzzy_results[0]['fuzzy_score']})")
        
        return fuzzy_results
    
    except ImportError:
        print("⚠️ fuzzywuzzy non installato, uso ricerca base")
        # Fallback a ricerca base
        all_docs = list_all_documents()
        results = [
            doc for doc in all_docs 
            if query.lower() in doc["name"].lower()
        ]
        return results
    
    except Exception as e:
        print(f"❌ Errore ricerca documenti: {e}")
        return []


# ============================
# CHUNKING - NUOVE FUNZIONI
# ============================

def read_document_in_chunks(file_path):
    """
    Legge un documento e lo divide in chunk.
    
    Returns:
        list: Lista di chunk con metadata
    """
    try:
        if not os.path.exists(file_path):
            print(f"❌ File non trovato: {file_path}")
            return []
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Leggi contenuto
        if file_ext == '.docx':
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            paragraphs = [p for p in content.split('\n') if p.strip()]
        else:
            return []
        
        # Dividi in chunk
        chunks = []
        for i in range(0, len(paragraphs), CHUNK_SIZE - CHUNK_OVERLAP):
            chunk_paragraphs = paragraphs[i:i + CHUNK_SIZE]
            chunk_text = '\n'.join(chunk_paragraphs)
            
            chunks.append({
                "text": chunk_text,
                "chunk_id": len(chunks),
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "relevance": 0
            })
        
        print(f"📄 Diviso {os.path.basename(file_path)} in {len(chunks)} chunk")
        return chunks
    
    except Exception as e:
        print(f"❌ Errore: {e}")
        return []


def search_in_chunks(query, max_chunks=3):
    """
    Cerca nei chunk e restituisce i più rilevanti.
    
    Args:
        query (str): Parola chiave
        max_chunks (int): Quanti chunk restituire
    
    Returns:
        list: Chunk rilevanti
    """
    try:
        all_docs = list_all_documents()
        relevant_chunks = []
        
        for doc in all_docs:
            chunks = read_document_in_chunks(doc["path"])
            
            for chunk in chunks:
                if query.lower() in chunk["text"].lower():
                    # Conta quante volte appare
                    relevance = chunk["text"].lower().count(query.lower())
                    chunk["relevance"] = relevance
                    relevant_chunks.append(chunk)
        
        # Ordina per rilevanza
        relevant_chunks.sort(key=lambda x: x["relevance"], reverse=True)
        top_chunks = relevant_chunks[:max_chunks]
        
        print(f"🔍 Trovati {len(top_chunks)} chunk rilevanti (da {len(relevant_chunks)} totali)")
        return top_chunks
    
    except Exception as e:
        print(f"❌ Errore: {e}")
        return []


def read_document(file_path, max_paragraphs=50):
    """
    Legge documento (LIMITATO a max_paragraphs per evitare overflow).
    Per file grandi, usa search_in_content invece!
    """
    try:
        if not os.path.exists(file_path):
            return ""
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            doc = Document(file_path)
            total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
            
            # Limita a primi N paragrafi
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()][:max_paragraphs]
            content = '\n'.join(paragraphs)
            
            # Avvisa se file è troncato
            if total_paragraphs > max_paragraphs:
                content += f"\n\n[⚠️ NOTA: File troppo lungo! Mostrati solo primi {max_paragraphs} paragrafi su {total_paragraphs}. Per cercare informazioni specifiche, usa la ricerca nel contenuto.]"
            
            print(f"📄 Letto {os.path.basename(file_path)} ({len(paragraphs)} paragrafi)")
        
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()[:max_paragraphs * 3]  # ~3 righe per paragrafo
                content = ''.join(lines)
            
            print(f"📄 Letto {os.path.basename(file_path)} (primi {len(lines)} righe)")
        
        else:
            return f"Formato non supportato: {file_ext}"
        
        return content
    
    except Exception as e:
        print(f"❌ Errore: {e}")
        return ""


def search_in_content(query):
    """
    Ricerca usando chunking (ottimizzato per file grandi).
    """
    try:
        # Cerca nei chunk
        relevant_chunks = search_in_chunks(query, max_chunks=5)
        
        if not relevant_chunks:
            return []
        
        results = []
        for chunk in relevant_chunks:
            lines = chunk["text"].split('\n')
            
            # Trova linee rilevanti (con la query)
            relevant_lines = [l for l in lines if query.lower() in l.lower()]
            
            # Se non trova linee specifiche, prendi le prime
            if not relevant_lines:
                relevant_lines = lines[:10]
            else:
                # Limita a max 10 linee rilevanti
                relevant_lines = relevant_lines[:10]
            
            # IMPORTANTE: NON includere full_content!
            # Claude riceve solo excerpt (10 righe invece di 5000 token)
            results.append({
                "name": chunk["file_name"],
                "path": chunk["file_path"],
                "chunk_id": chunk["chunk_id"],
                "excerpt": relevant_lines,  # Solo 10 righe (~500 token)
                "relevance": chunk.get("relevance", 0)
            })
        
        print(f"🔍 {len(results)} risultati (ottimizzato per token)")
        return results
    
    except Exception as e:
        print(f"❌ Errore: {e}")
        return []

# ============================
# TEST
# ============================

def test_reader():
    print("\n" + "="*50)
    print("TEST CHUNKING")
    print("="*50)
    
    docs = list_all_documents()
    print(f"\n📋 {len(docs)} documenti")
    
    if docs:
        chunks = read_document_in_chunks(docs[0]["path"])
        print(f"\n🧩 {len(chunks)} chunk creati")
        if chunks:
            print(f"  Primo chunk: {chunks[0]['text'][:200]}...")
    
    print("\n🔍 Test ricerca chunk...")
    results = search_in_chunks("prescrizione", max_chunks=2)
    print(f"  {len(results)} chunk rilevanti")
    
    print("\n" + "="*50)


if __name__ == "__main__":
    test_reader()